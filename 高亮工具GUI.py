# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import os
import regex as re
import logging
import sys
import threading
import tempfile
import zipfile
from copy import deepcopy
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx.opc.exceptions import PackageNotFoundError
from docx.text.run import Run

APP_NAME = "Word高亮工具"

def get_app_dir():
    """获取程序运行时的根目录，兼容 PyInstaller 打包后的路径"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def ensure_writable_dir(path):
    """创建并确认目录可写。"""
    os.makedirs(path, exist_ok=True)
    probe_path = os.path.join(path, ".write_test")
    with open(probe_path, "w", encoding="utf-8") as f:
        f.write("")
    os.remove(probe_path)
    return path

def get_user_data_dir():
    """获取日志和默认关键词配置使用的可写目录。"""
    if not getattr(sys, 'frozen', False):
        return get_app_dir()

    candidates = []
    if sys.platform == "darwin":
        candidates.extend([
            os.path.join(os.path.expanduser("~/Documents"), APP_NAME),
            os.path.join(os.path.expanduser("~/Library/Application Support"), APP_NAME),
        ])
    else:
        candidates.extend([
            get_app_dir(),
            os.path.join(os.path.expanduser("~/Documents"), APP_NAME),
        ])

    candidates.append(os.path.join(tempfile.gettempdir(), APP_NAME))

    for candidate in candidates:
        try:
            return ensure_writable_dir(candidate)
        except Exception:
            continue

    return tempfile.gettempdir()

def get_log_file_path():
    return os.path.join(get_user_data_dir(), "highlight_gui.log")

# --- 日志配置 ---
# 创建一个基础的 Logger
logger = logging.getLogger("HighlightApp")
logger.setLevel(logging.INFO)

# 默认控制台和文件日志
log_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
log_file_handler = logging.FileHandler(get_log_file_path(), 'w', encoding='utf-8')
log_file_handler.setFormatter(log_formatter)
logger.addHandler(log_file_handler)

def log_uncaught_exception(exc_type, exc_value, exc_traceback):
    if issubclass(exc_type, KeyboardInterrupt):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logger.critical("程序发生未捕获异常", exc_info=(exc_type, exc_value, exc_traceback))

sys.excepthook = log_uncaught_exception

def ensure_default_keywords_files():
    """自动创建默认关键词文件夹及空的关键词文本文件"""
    app_dir = get_user_data_dir()
    default_dir = os.path.join(app_dir, "关键词配置")
    if not os.path.exists(default_dir):
        try:
            os.makedirs(default_dir)
            logger.info(f"自动创建默认关键词文件夹: {default_dir}")
        except Exception as e:
            logger.error(f"创建默认关键词文件夹失败: {e}")
            return default_dir
            
    files = ["高亮关键词.txt", "特殊关键词.txt", "排除关键词.txt"]
    for file_name in files:
        file_path = os.path.join(default_dir, file_name)
        if not os.path.exists(file_path):
            try:
                with open(file_path, 'w', encoding='utf-8-sig') as f:
                    f.write('') # 创建空文件
                logger.info(f"自动创建默认关键词文件: {file_name}")
            except Exception as e:
                logger.error(f"创建默认关键词文件 {file_name} 失败: {e}")
    return default_dir

# --- 核心高亮逻辑函数 ---
def load_keywords(filepath):
    """从文本文件加载关键词列表，去除空行和首尾空格。"""
    logger.info(f"开始加载关键词文件: {os.path.basename(filepath)}")
    if not os.path.exists(filepath):
        logger.warning(f"警告: 关键词文件未找到 {os.path.basename(filepath)}，将使用空列表。")
        return []
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            keywords = [line.strip() for line in f if line.strip()]
        logger.info(f"成功加载 {len(keywords)} 个关键词 (来自 {os.path.basename(filepath)})")
        return keywords
    except Exception as e:
        logger.error(f"加载关键词文件 {os.path.basename(filepath)} 时出错: {e}")
        return []

def get_docx_validation_error(filepath):
    """返回 Word 文档校验错误；有效时返回 None。"""
    if not os.path.exists(filepath):
        return f"找不到指定的 Word 文件：\n{filepath}"

    try:
        if os.path.getsize(filepath) == 0:
            return f"该 Word 文件大小为 0 字节，无法读取：\n{filepath}"

        if not zipfile.is_zipfile(filepath):
            return f"该文件不是有效的 .docx 文件，请选择真正的 Word 文档：\n{filepath}"

        with zipfile.ZipFile(filepath) as zf:
            names = set(zf.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                return f"该文件不是标准 Word .docx 文档，无法处理：\n{filepath}"
    except OSError as e:
        return f"无法读取该 Word 文件：\n{filepath}\n\n{e}"
    except zipfile.BadZipFile:
        return f"该文件损坏或不是有效的 .docx 文件：\n{filepath}"

    return None

def append_cloned_run(paragraph, source_run, text=None, highlight_color=None):
    """克隆原始 run 的完整 XML 格式，只替换文本并按需设置高亮。"""
    new_r = deepcopy(source_run._r)
    new_run = Run(new_r, paragraph)
    if text is not None:
        new_run.text = text
    paragraph._p.append(new_r)
    if highlight_color is not None:
        new_run.font.highlight_color = highlight_color
    return new_run

def rebuild_paragraph_with_highlights(paragraph, original_runs, merged_spans, highlight_color):
    """按高亮区间重建段落，同时完整保留每个原始 run 的格式。"""
    paragraph.clear()

    char_pos = 0
    span_index = 0

    for source_run in original_runs:
        run_text = source_run.text
        run_len = len(run_text)

        if run_len == 0:
            append_cloned_run(paragraph, source_run)
            continue

        run_end = char_pos + run_len
        offset = 0

        while offset < run_len:
            abs_pos = char_pos + offset

            while span_index < len(merged_spans) and merged_spans[span_index][1] <= abs_pos:
                span_index += 1

            is_highlight = (
                span_index < len(merged_spans)
                and merged_spans[span_index][0] <= abs_pos < merged_spans[span_index][1]
            )

            if is_highlight:
                next_abs = min(run_end, merged_spans[span_index][1])
            else:
                next_abs = run_end
                if span_index < len(merged_spans):
                    next_abs = min(next_abs, merged_spans[span_index][0])

            if next_abs <= abs_pos:
                next_abs = abs_pos + 1

            part = run_text[offset:offset + (next_abs - abs_pos)]
            append_cloned_run(
                paragraph,
                source_run,
                part,
                highlight_color if is_highlight else None
            )
            offset += next_abs - abs_pos

        char_pos = run_end

def has_cjk(text):
    """检查字符串是否包含中日韩字符。"""
    return any('\u4e00' <= char <= '\u9fff' for char in text)

def create_pattern_for_keyword(kw, match_type):
    """根据关键词和匹配类型创建智能的正则表达式。"""
    escaped_kw = re.escape(kw)
    
    if has_cjk(kw):
        return escaped_kw

    if match_type == 'highlight':
        if ' ' in kw:
            return f'\\b{escaped_kw}\\b'
        else:
            return escaped_kw
            
    elif match_type in ['exact', 'exclude']:
        return f'\\b{escaped_kw}\\b'
    
    return escaped_kw

def process_and_highlight_paragraph(paragraph, highlight_keywords, exact_keywords, exclude_keywords, highlight_color, case_sensitive=False):
    """处理单个段落的核心逻辑。"""
    if not paragraph.text.strip():
        return

    paragraph_text = paragraph.text
    potential_spans = []
    regex_flags = 0 if case_sensitive else re.IGNORECASE

    # 1a: 处理 "高亮关键词.txt" (包含匹配)
    for kw in highlight_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'highlight')
            for match in re.finditer(pattern, paragraph_text, regex_flags):
                start, end = match.span()
                potential_spans.append((start, end, kw))
        except re.error as e:
            logger.error(f"  正则表达式错误 (高亮)，关键词: '{kw}', 错误: {e}")

    # 1b: 处理 "特殊关键词.txt" (绝对匹配)
    for kw in exact_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'exact')
            for match in re.finditer(pattern, paragraph_text, regex_flags):
                potential_spans.append((match.start(), match.end(), kw))
        except re.error as e:
            logger.error(f"  正则表达式错误 (特殊)，关键词: '{kw}', 错误: {e}")
    
    if not potential_spans:
        return

    # 步骤 2: 查找所有需要排除的区域
    exclude_spans = []
    for kw in exclude_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'exclude')
            for match in re.finditer(pattern, paragraph_text, regex_flags):
                exclude_spans.append((match.start(), match.end(), kw))
        except re.error as e:
            logger.error(f"  正则表达式错误 (排除)，关键词: '{kw}', 错误: {e}")

    # 步骤 3: 应用排除规则
    final_spans = []
    for p_start, p_end, p_kw in potential_spans:
        is_excluded = False
        for e_start, e_end, e_kw in exclude_spans:
            if p_start >= e_start and p_end <= e_end:
                is_excluded = True
                break
        if not is_excluded:
            final_spans.append((p_start, p_end))

    if not final_spans:
        return

    # 步骤 4: 合并重叠或相邻的高亮区域
    final_spans.sort(key=lambda x: x[0])
    merged_spans = []
    current_start, current_end = final_spans[0]
    for next_start, next_end in final_spans[1:]:
        if next_start < current_end:
            current_end = max(current_end, next_end)
        else:
            merged_spans.append((current_start, current_end))
            current_start, current_end = next_start, next_end
    merged_spans.append((current_start, current_end))

    original_runs = list(paragraph.runs)
    rebuild_paragraph_with_highlights(paragraph, original_runs, merged_spans, highlight_color)


# --- GUI 桥接日志 Handler ---
class TextboxHandler(logging.Handler):
    """自定义日志 Handler，用于将 Python logging 的日志输出实时写入 CustomTkinter Textbox"""
    def __init__(self, textbox):
        super().__init__()
        self.textbox = textbox

    def emit(self, record):
        msg = self.format(record)
        # 线程安全更新 GUI
        self.textbox.after(0, self.append_text, msg + "\n")

    def append_text(self, text):
        self.textbox.configure(state="normal")
        self.textbox.insert("end", text)
        self.textbox.see("end")
        self.textbox.configure(state="disabled")


# --- GUI 主类 ---
class WordHighlighterApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # 窗口基本属性配置
        self.title("Word 关键词高亮智能工具")
        self.geometry("760x650")
        self.minsize(700, 580)
        
        # 居中显示
        self.center_window()

        # 配置默认主题色
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        # 帮助窗口引用
        self.help_window = None

        # 颜色映射字典
        self.color_map = {
            "黄色 (Yellow)": WD_COLOR_INDEX.YELLOW,
            "绿色 (Green)": WD_COLOR_INDEX.GREEN,
            "青色 (Turquoise)": WD_COLOR_INDEX.TURQUOISE,
            "粉色 (Pink)": WD_COLOR_INDEX.PINK,
            "红色 (Red)": WD_COLOR_INDEX.RED,
            "蓝色 (Blue)": WD_COLOR_INDEX.BLUE,
        }

        # 建立网格权重
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(3, weight=1) # 让日志文本框所在的行拉伸

        # --- 1. 标题与设置栏 ---
        self.header_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.header_frame.grid(row=0, column=0, padx=20, pady=(15, 10), sticky="ew")
        self.header_frame.grid_columnconfigure(0, weight=1)
        
        self.title_label = ctk.CTkLabel(
            self.header_frame, 
            text="Word 关键词高亮智能工具", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=22, weight="bold")
        )
        self.title_label.grid(row=0, column=0, sticky="w")
        
        # 使用说明按钮
        self.help_button = ctk.CTkButton(
            self.header_frame,
            text="❓ 使用说明",
            width=90,
            height=28,
            fg_color="transparent",
            text_color=("#1f77b4", "#52a3db"),
            hover_color=("#e6f2fc", "#2a3c4d"),
            command=self.open_help_window,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12, weight="bold")
        )
        self.help_button.grid(row=0, column=1, padx=(10, 20), sticky="e")
        
        # 主题切换开关
        self.theme_switch = ctk.CTkSwitch(
            self.header_frame, 
            text="深色模式", 
            command=self.toggle_theme,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.theme_switch.grid(row=0, column=2, sticky="e")
        if ctk.get_appearance_mode() == "Dark":
            self.theme_switch.select()

        # --- 2. 文件与文件夹选择框卡片 ---
        self.config_frame = ctk.CTkFrame(self, border_width=1, corner_radius=12)
        self.config_frame.grid(row=1, column=0, padx=20, pady=10, sticky="ew")
        self.config_frame.grid_columnconfigure(1, weight=1)

        # 2a. Word 文档路径选择
        self.doc_label = ctk.CTkLabel(
            self.config_frame, 
            text="选择 Word 文档:", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        self.doc_label.grid(row=0, column=0, padx=(15, 10), pady=(15, 10), sticky="w")
        
        self.doc_entry = ctk.CTkEntry(
            self.config_frame, 
            placeholder_text="点击右侧按钮选择需要高亮的 .docx 文件...",
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.doc_entry.grid(row=0, column=1, padx=10, pady=(15, 10), sticky="ew")
        # 设为只读以防止误输入，但由于 Tkinter 机制，我们使用自定义行为，或者用 CTkEntry 但让用户只能通过按钮选择
        self.doc_entry.configure(state="readonly")
        
        self.doc_button = ctk.CTkButton(
            self.config_frame, 
            text="选择文件", 
            width=100,
            command=self.select_docx_file,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.doc_button.grid(row=0, column=2, padx=(10, 15), pady=(15, 10))

        # 2b. 关键词文件夹选择
        self.folder_label = ctk.CTkLabel(
            self.config_frame, 
            text="关键词文件夹:", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        self.folder_label.grid(row=1, column=0, padx=(15, 10), pady=10, sticky="w")
        
        self.folder_entry = ctk.CTkEntry(
            self.config_frame, 
            placeholder_text="选择存放高亮/特殊/排除关键词 txt 文件的文件夹...",
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.folder_entry.grid(row=1, column=1, padx=10, pady=10, sticky="ew")
        
        # 自动初始化默认文件夹和空关键词文件并填充
        default_dir = ensure_default_keywords_files()
        self.folder_entry.insert(0, default_dir)
        self.folder_entry.configure(state="readonly")
        
        self.folder_button = ctk.CTkButton(
            self.config_frame, 
            text="选择文件夹", 
            width=100,
            command=self.select_keywords_folder,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.folder_button.grid(row=1, column=2, padx=(10, 15), pady=10)

        # 2c. 高亮颜色选择与提示语
        self.color_label = ctk.CTkLabel(
            self.config_frame, 
            text="设定高亮颜色:", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        self.color_label.grid(row=2, column=0, padx=(15, 10), pady=(10, 15), sticky="w")

        self.color_menu = ctk.CTkOptionMenu(
            self.config_frame, 
            values=list(self.color_map.keys()),
            width=160,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.color_menu.grid(row=2, column=1, padx=10, pady=(10, 15), sticky="w")
        self.color_menu.set("黄色 (Yellow)")

        self.info_label = ctk.CTkLabel(
            self.config_frame, 
            text="提示：请确保所选文件夹内包含 `高亮关键词.txt` 等文件", 
            text_color="gray",
            font=ctk.CTkFont(family="Microsoft YaHei", size=11)
        )
        self.info_label.grid(row=2, column=1, columnspan=2, padx=(180, 15), pady=(10, 15), sticky="e")

        # 2d. 大小写匹配开关
        self.case_label = ctk.CTkLabel(
            self.config_frame,
            text="大小写匹配:",
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        self.case_label.grid(row=3, column=0, padx=(15, 10), pady=(0, 15), sticky="w")

        self.case_switch = ctk.CTkSwitch(
            self.config_frame,
            text="区分大小写",
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.case_switch.grid(row=3, column=1, padx=10, pady=(0, 15), sticky="w")

        self.case_hint_label = ctk.CTkLabel(
            self.config_frame,
            text="关闭时不区分大小写（默认）",
            text_color="gray",
            font=ctk.CTkFont(family="Microsoft YaHei", size=11)
        )
        self.case_hint_label.grid(row=3, column=1, columnspan=2, padx=(130, 15), pady=(0, 15), sticky="e")

        # --- 3. 进度条与控制区 ---
        self.control_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.control_frame.grid(row=2, column=0, padx=20, pady=10, sticky="ew")
        self.control_frame.grid_columnconfigure(0, weight=1)

        self.progress_label = ctk.CTkLabel(
            self.control_frame, 
            text="等待开始...", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=12)
        )
        self.progress_label.grid(row=0, column=0, padx=5, pady=(0, 5), sticky="w")

        self.progress_bar = ctk.CTkProgressBar(self.control_frame, height=12, corner_radius=6)
        self.progress_bar.grid(row=1, column=0, padx=5, pady=5, sticky="ew")
        self.progress_bar.set(0.0)

        self.run_button = ctk.CTkButton(
            self.control_frame, 
            text="开始高亮处理", 
            height=40,
            width=160,
            fg_color="#1f77b4",
            hover_color="#105a8f",
            command=self.start_processing_thread,
            font=ctk.CTkFont(family="Microsoft YaHei", size=14, weight="bold")
        )
        self.run_button.grid(row=0, column=1, rowspan=2, padx=(20, 5), pady=5, sticky="ns")

        # --- 4. 运行日志区 ---
        self.log_label = ctk.CTkLabel(
            self, 
            text="运行日志:", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        self.log_label.grid(row=3, column=0, padx=25, pady=(10, 0), sticky="w")

        self.log_textbox = ctk.CTkTextbox(
            self, 
            corner_radius=8,
            font=ctk.CTkFont(family="Consolas", size=11),
            wrap="word",
            border_width=1
        )
        self.log_textbox.grid(row=4, column=0, padx=20, pady=(5, 20), sticky="nsew")
        self.log_textbox.configure(state="disabled")

        # 将日志绑定到 TextBox
        self.textbox_handler = TextboxHandler(self.log_textbox)
        self.textbox_handler.setFormatter(log_formatter)
        logger.addHandler(self.textbox_handler)

    def center_window(self):
        """让窗口居中于屏幕"""
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f'+{x}+{y}')

    def toggle_theme(self):
        """切换深色/浅色主题"""
        if self.theme_switch.get() == 1:
            ctk.set_appearance_mode("Dark")
        else:
            ctk.set_appearance_mode("Light")

    def open_help_window(self):
        """打开/显示使用说明弹窗 (单例模式)"""
        if hasattr(self, "help_window") and self.help_window is not None and self.help_window.winfo_exists():
            self.help_window.deiconify()
            self.help_window.focus()
            return

        self.help_window = ctk.CTkToplevel(self)
        self.help_window.title("使用说明与步骤")
        self.help_window.geometry("620x420")
        self.help_window.resizable(False, False)
        
        self.help_window.update_idletasks()
        width = 620
        height = 420
        x = self.winfo_x() + (self.winfo_width() // 2) - (width // 2)
        y = self.winfo_y() + (self.winfo_height() // 2) - (height // 2)
        self.help_window.geometry(f"{width}x{height}+{x}+{y}")
        
        self.help_window.transient(self)
        self.help_window.grab_set()

        self.help_window.grid_columnconfigure(0, weight=1)
        self.help_window.grid_rowconfigure(0, weight=1)

        content_frame = ctk.CTkFrame(self.help_window, corner_radius=12, border_width=1)
        content_frame.grid(row=0, column=0, padx=15, pady=15, sticky="nsew")
        content_frame.grid_columnconfigure(0, weight=1)

        help_title = ctk.CTkLabel(
            content_frame, 
            text="📖 Word 关键词高亮工具使用说明", 
            font=ctk.CTkFont(family="Microsoft YaHei", size=16, weight="bold")
        )
        help_title.grid(row=0, column=0, padx=15, pady=(15, 10), sticky="w")

        help_text_box = ctk.CTkTextbox(
            content_frame,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12),
            wrap="word",
            fg_color="transparent",
            activate_scrollbars=True
        )
        help_text_box.grid(row=1, column=0, padx=15, pady=5, sticky="nsew")
        content_frame.grid_rowconfigure(1, weight=1)

        help_content = (
            "程序启动时，会在程序所在路径下自动创建一个【关键词配置】文件夹，"
            "并在该文件夹下生成以下三个 .txt 规则文件。你只需在相应的文件中填入关键词即可：\n\n"
            "1. 💡 高亮关键词.txt (包含匹配/模糊匹配)\n"
            "   - 用于放置需要高亮的关键词。如果是中文，则进行包含匹配；如果是英文单词，则会匹配包含它的内容（不限制边界）。\n\n"
            "2. 🎯 特殊关键词.txt (绝对/全词匹配)\n"
            "   - 用于进行绝对/全词匹配。例如，希望仅在英文单词独立出现时高亮（如只匹配 apple 而不匹配 apples），请将其放入此文件。\n\n"
            "3. 🚫 排除关键词.txt (排除标记)\n"
            "   - 用于防止误标记。例如，希望高亮 AI，但当出现 AI Agent 时不希望高亮其中的 AI。此时可以在 排除关键词.txt 中填入 AI Agent。\n"
            "\n4. 🔠 大小写匹配\n"
            "   - 默认关闭：不区分大小写；开启后：只匹配大小写完全一致的内容。\n"
        )
        
        help_text_box.insert("1.0", help_content)
        help_text_box.configure(state="disabled")

        warning_label = ctk.CTkLabel(
            content_frame,
            text="⚠️ 注意：每个文本文件中一行只能放置一个关键词！",
            text_color=("#d9534f", "#ff4d4d"),
            font=ctk.CTkFont(family="Microsoft YaHei", size=13, weight="bold")
        )
        warning_label.grid(row=2, column=0, padx=15, pady=(10, 5), sticky="w")

        close_btn = ctk.CTkButton(
            content_frame,
            text="我知道了",
            width=100,
            command=self.help_window.destroy,
            font=ctk.CTkFont(family="Microsoft YaHei", size=12, weight="bold")
        )
        close_btn.grid(row=3, column=0, padx=15, pady=(5, 15))

    def select_docx_file(self):
        """选择 Word 文档并更新 Entry"""
        # 如果当前在 Windows 上，默认打开桌面
        init_dir = os.path.expanduser("~/Desktop")
        file_path = filedialog.askopenfilename(
            initialdir=init_dir,
            title="选择要高亮的 Word 文档",
            filetypes=[("Word 文档", "*.docx")]
        )
        if file_path:
            # 更新 Entry 内容 (由于是 readonly 状态，需要临时改为 normal 写入后再改回)
            self.doc_entry.configure(state="normal")
            self.doc_entry.delete(0, "end")
            self.doc_entry.insert(0, file_path)
            self.doc_entry.configure(state="readonly")
            
            # 如果关键词文件夹还没有选择，自动推导为该 Word 文档所在目录
            if not self.folder_entry.get():
                folder_path = os.path.dirname(file_path)
                self.folder_entry.configure(state="normal")
                self.folder_entry.delete(0, "end")
                self.folder_entry.insert(0, folder_path)
                self.folder_entry.configure(state="readonly")

    def select_keywords_folder(self):
        """选择存放关键词文本的文件夹并更新 Entry"""
        init_dir = os.path.dirname(self.doc_entry.get()) if self.doc_entry.get() else os.path.expanduser("~/Desktop")
        folder_path = filedialog.askdirectory(
            initialdir=init_dir,
            title="选择存放关键词 txt 文件的文件夹"
        )
        if folder_path:
            self.folder_entry.configure(state="normal")
            self.folder_entry.delete(0, "end")
            self.folder_entry.insert(0, folder_path)
            self.folder_entry.configure(state="readonly")

    def set_widgets_state(self, state):
        """启用或禁用界面交互控件"""
        self.doc_button.configure(state=state)
        self.folder_button.configure(state=state)
        self.color_menu.configure(state=state)
        self.case_switch.configure(state=state)
        self.run_button.configure(state=state)
        self.theme_switch.configure(state=state)

    def safe_update_progress(self, current, total, text_desc):
        """线程安全地更新进度条和描述"""
        val = current / total if total > 0 else 0
        self.after(0, lambda: self.progress_bar.set(val))
        self.after(0, lambda: self.progress_label.configure(text=f"{text_desc} ({int(val * 100)}%)"))

    def start_processing_thread(self):
        """获取参数，执行参数校验，启动后台工作线程"""
        doc_path = self.doc_entry.get().strip()
        folder_path = self.folder_entry.get().strip()
        selected_color_name = self.color_menu.get()
        case_sensitive = self.case_switch.get() == 1

        if not doc_path:
            messagebox.showerror("参数缺失", "请先选择需要高亮的 Word 文档！")
            return

        docx_error = get_docx_validation_error(doc_path)
        if docx_error:
            messagebox.showerror("Word 文档无效", docx_error)
            return

        if not folder_path:
            messagebox.showerror("参数缺失", "请选择存放关键词 txt 文件的文件夹！")
            return
        if not os.path.exists(folder_path):
            messagebox.showerror("目录不存在", f"找不到指定的文件夹：\n{folder_path}")
            return

        # 检查文件夹下是否至少有高亮关键词.txt
        hl_txt = os.path.join(folder_path, "高亮关键词.txt")
        if not os.path.exists(hl_txt):
            # 询问用户是否继续
            confirm = messagebox.askyesno(
                "关键词文件缺失", 
                f"在文件夹中未找到核心的 `高亮关键词.txt`。\n是否仍要继续处理（如果没有有效关键词，文档将不会进行任何高亮操作）？"
            )
            if not confirm:
                return

        # 获取对应的高亮颜色枚举值
        highlight_color = self.color_map.get(selected_color_name, WD_COLOR_INDEX.YELLOW)

        # 禁用按钮，防重复点击
        self.set_widgets_state("disabled")
        self.progress_bar.set(0.0)
        self.progress_label.configure(text="开始处理中...")
        
        # 清空界面日志
        self.log_textbox.configure(state="normal")
        self.log_textbox.delete("1.0", "end")
        self.log_textbox.configure(state="disabled")

        # 开启后台工作线程
        worker = threading.Thread(
            target=self.run_highlighting_task,
            args=(doc_path, folder_path, highlight_color, case_sensitive)
        )
        worker.daemon = True
        worker.start()

    def run_highlighting_task(self, doc_path, folder_path, highlight_color, case_sensitive):
        """在后台线程中执行高亮处理逻辑"""
        try:
            # 构造输入文件路径
            highlight_keywords_path = os.path.join(folder_path, '高亮关键词.txt')
            exact_keywords_path = os.path.join(folder_path, '特殊关键词.txt')
            exclude_keywords_path = os.path.join(folder_path, '排除关键词.txt')

            # 构造输出文件路径：高亮后的文件放到高亮文档旁边，多增加2个字 高亮
            dir_name, file_name = os.path.split(doc_path)
            base_name, ext = os.path.splitext(file_name)
            output_path = os.path.join(dir_name, f"{base_name}高亮{ext}")

            logger.info("--- 开启关键词高亮任务 ---")
            logger.info(f"大小写匹配模式: {'区分大小写' if case_sensitive else '不区分大小写'}")

            docx_error = get_docx_validation_error(doc_path)
            if docx_error:
                logger.error(docx_error)
                self.after(0, lambda: self.on_task_finished(False, docx_error))
                return
            
            # 1. 加载关键词
            logger.info("正在加载关键词文本...")
            highlight_keywords = load_keywords(highlight_keywords_path)
            exact_keywords = load_keywords(exact_keywords_path)
            exclude_keywords = load_keywords(exclude_keywords_path)

            if not highlight_keywords and not exact_keywords:
                logger.warning("警告: '高亮关键词' 和 '特殊关键词' 列表都为空，将不会进行任何高亮标记。")

            # 2. 打开 Word
            logger.info(f"正在加载 Word 文档: {os.path.basename(doc_path)}")
            document = docx.Document(doc_path)
            logger.info("Word 文档加载成功。")

            # 3. 统计总段落数和表格数以计算进度
            total_paras = len(document.paragraphs)
            total_tables = len(document.tables)
            total_steps = total_paras + total_tables

            if total_steps == 0:
                logger.warning("该 Word 文档没有任何段落或表格，无需处理。")
                self.after(0, lambda: self.on_task_finished(True, output_path))
                return

            current_step = 0

            # 4. 处理正文段落
            logger.info("开始高亮主文档段落...")
            for i, p in enumerate(document.paragraphs):
                process_and_highlight_paragraph(
                    p, highlight_keywords, exact_keywords, exclude_keywords, highlight_color, case_sensitive
                )
                current_step += 1
                if current_step % 10 == 0 or current_step == total_steps:
                    self.safe_update_progress(current_step, total_steps, f"正在处理正文段落 {i+1}/{total_paras}")
            
            # 5. 处理表格内容
            logger.info("开始高亮表格段落...")
            for table_idx, table in enumerate(document.tables):
                for row in table.rows:
                    for cell in row.cells:
                        for p_in_cell in cell.paragraphs:
                            process_and_highlight_paragraph(
                                p_in_cell, highlight_keywords, exact_keywords, exclude_keywords, highlight_color, case_sensitive
                            )
                current_step += 1
                self.safe_update_progress(current_step, total_steps, f"正在处理表格 {table_idx+1}/{total_tables}")

            # 6. 保存新文档
            logger.info(f"正在保存标记高亮后的文档到原文档同目录...")
            document.save(output_path)
            logger.info(f"保存成功！新文件：{os.path.basename(output_path)}")
            logger.info("--- 任务处理成功结束 ---")

            self.after(0, lambda: self.on_task_finished(True, output_path))

        except PermissionError:
            err_msg = f"保存文档失败！输出文件 '{os.path.basename(output_path)}' 可能已被其他程序（如Word）打开，请先关闭该文件并重试。"
            logger.error(f"错误: {err_msg}")
            self.after(0, lambda: self.on_task_finished(False, err_msg))
        except (PackageNotFoundError, zipfile.BadZipFile):
            err_msg = f"Word 文档无效或已损坏，无法读取：\n{doc_path}\n\n请确认选择的是原始 .docx 文件，而不是 0 字节文件或未生成成功的输出文件。"
            logger.error(f"错误: {err_msg}")
            self.after(0, lambda: self.on_task_finished(False, err_msg))
        except Exception as e:
            err_msg = f"处理过程中发生未知错误: {str(e)}"
            logger.error(err_msg, exc_info=True)
            self.after(0, lambda: self.on_task_finished(False, err_msg))

    def on_task_finished(self, success, result_msg):
        """当后台高亮处理任务结束时的回调"""
        self.set_widgets_state("normal")
        if success:
            self.progress_bar.set(1.0)
            self.progress_label.configure(text="处理完成！")
            messagebox.showinfo("处理成功", f"文件高亮标记已处理完成！\n\n新文件已生成在：\n{result_msg}")
        else:
            self.progress_bar.set(0.0)
            self.progress_label.configure(text="处理失败")
            messagebox.showerror("运行出错", result_msg)


def main():
    # 启用 DPI 缩放（在高分屏/视网膜屏上防止界面模糊）
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    logger.info(f"程序启动，日志文件: {get_log_file_path()}")
    logger.info(f"默认数据目录: {get_user_data_dir()}")

    app = WordHighlighterApp()
    app.mainloop()


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logger.critical("程序启动失败", exc_info=True)
        try:
            messagebox.showerror(
                "启动失败",
                f"程序启动失败：{e}\n\n请查看日志文件：\n{get_log_file_path()}"
            )
        except Exception:
            pass
        sys.exit(1)
