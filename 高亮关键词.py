# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import os
import regex as re # 使用更强大的 regex 库
import logging
import sys

# --- 日志配置 ---
log_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
log_file_handler = logging.FileHandler('highlight_script.log', 'w', encoding='utf-8')
log_file_handler.setFormatter(log_formatter)
log_console_handler = logging.StreamHandler(sys.stdout)
log_console_handler.setFormatter(log_formatter)

logger = logging.getLogger()
logger.setLevel(logging.INFO) # 设置日志级别 INFO 或 DEBUG
if logger.hasHandlers():
    logger.handlers.clear()
logger.addHandler(log_file_handler)
logger.addHandler(log_console_handler)
# --- 日志配置结束 ---


# --- 配置 ---
WORD_DOCUMENT_PATH = r'C:\Users\newnew\Desktop\高亮标记\00【底稿】《卷七 关于追求真理》最新定稿_20251228.docx'
HIGHLIGHT_KEYWORDS_PATH = os.path.join('关键词配置', '高亮关键词.txt')      # 规则 1: 包含匹配 (只高亮关键词部分)
EXACT_KEYWORDS_PATH = os.path.join('关键词配置', '特殊关键词.txt')          # 规则 3: 绝对匹配 (整个单词/短语)
EXCLUDE_KEYWORDS_PATH = os.path.join('关键词配置', '排除关键词.txt')      # 规则 2: 排除匹配
OUTPUT_DOCUMENT_PATH = r'C:\Users\newnew\Desktop\高亮标记\标记高亮_00【底稿】《卷七 关于追求真理》最新定稿_20251228.docx' # 输出文件名
HIGHLIGHT_COLOR = WD_COLOR_INDEX.YELLOW       # 高亮颜色
# --- 配置结束 ---


# --- 工具函数 ---
def get_app_dir():
    """获取程序运行时的根目录，兼容 PyInstaller 打包后的路径"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def ensure_default_keywords_files():
    """自动创建默认关键词文件夹及空的关键词文本文件"""
    app_dir = get_app_dir()
    default_dir = os.path.join(app_dir, "关键词配置")
    if not os.path.exists(default_dir):
        try:
            os.makedirs(default_dir)
            logger.info(f"自动创建默认关键词文件夹: {default_dir}")
        except Exception as e:
            logger.error(f"创建默认关键词文件夹失败: {e}")
            return default_dir
            
    files = ["高亮关键词.txt", "特殊关键词.txt", "排除关键词.txt"]
    for file_name in files:
        file_path = os.path.join(default_dir, file_name)
        if not os.path.exists(file_path):
            try:
                with open(file_path, 'w', encoding='utf-8-sig') as f:
                    f.write('') # 创建空文件
                logger.info(f"自动创建默认关键词文件: {file_name}")
            except Exception as e:
                logger.error(f"创建默认关键词文件 {file_name} 失败: {e}")
    return default_dir

def load_keywords(filepath):
    """从文本文件加载关键词列表，去除空行和首尾空格。"""
    logger.info(f"开始加载关键词文件: {filepath}")
    if not os.path.exists(filepath):
        logger.warning(f"警告: 关键词文件未找到 {filepath}，将使用空列表。")
        return []
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            keywords = [line.strip() for line in f if line.strip()]
        logger.info(f"成功加载 {len(keywords)} 个关键词 (来自 {filepath})")
        return keywords
    except Exception as e:
        logger.error(f"加载关键词文件 {filepath} 时出错: {e}")
        return []

def copy_run_format(source_run, target_run):
    """复制 run 的格式，包括字体、字号、颜色、粗体、斜体、以及原有的高亮颜色。"""
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.style.name = source_run.style.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    # --- 【代码修复】 ---
    # 复制原始 run 的高亮颜色，以保留原有的背景色
    if source_run.font.highlight_color:
        target_run.font.highlight_color = source_run.font.highlight_color
    # --- 【修复结束】 ---


def has_cjk(text):
    """检查字符串是否包含中日韩字符。"""
    # CJK Unicode 范围
    return any('\u4e00' <= char <= '\u9fff' for char in text)

def create_pattern_for_keyword(kw, match_type):
    """
    【核心改进】根据关键词和匹配类型创建智能的正则表达式。
    match_type: 'highlight', 'exact', 'exclude'
    """
    escaped_kw = re.escape(kw)
    
    # 如果是中文，不使用单词边界 \b
    if has_cjk(kw):
        return escaped_kw

    # 如果是英文或其它字母语言
    if match_type == 'highlight':
        # 高亮规则: 
        # 1. 如果是短语 (含空格)，则全词匹配，避免匹配到单词的一部分。
        # 2. 如果是单个词，则进行包含匹配 (不加 \b)，允许匹配单词的任何部分。
        if ' ' in kw:
            return f'\\b{escaped_kw}\\b'
        else:
            # --- 这是修复的关键 ---
            # 对于单个词的包含匹配，不应该使用 \b
            return escaped_kw
            
    elif match_type in ['exact', 'exclude']:
        # 精确匹配和排除规则：必须是全词匹配
        return f'\\b{escaped_kw}\\b'
    
    return escaped_kw # 默认返回转义后的关键词

# --- 核心处理逻辑 ---

def process_and_highlight_paragraph(paragraph, highlight_keywords, exact_keywords, exclude_keywords, highlight_color):
    """
    处理单个段落的核心逻辑。
    """
    if not paragraph.text.strip():
        return

    paragraph_text = paragraph.text
    logger.debug(f"\n== 处理段落 (首50字符): {paragraph_text[:50].replace(chr(10), ' ')} ==")

    # 步骤 1: 查找所有可能的匹配位置 (spans)
    potential_spans = []

    # 1a: 处理 "高亮关键词.txt" (包含匹配)
    for kw in highlight_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'highlight')
            # 使用 finditer 查找所有不区分大小写的匹配项
            for match in re.finditer(pattern, paragraph_text, re.IGNORECASE):
                start, end = match.span()
                potential_spans.append((start, end, kw))
                logger.debug(f"  找到[高亮]匹配: '{paragraph_text[start:end]}' (规则: '{kw}') at [{start}:{end}]")
        except re.error as e:
            logger.error(f"  正则表达式错误 (高亮)，关键词: '{kw}', 错误: {e}")

    # 1b: 处理 "特殊关键词.txt" (绝对匹配)
    for kw in exact_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'exact')
            for match in re.finditer(pattern, paragraph_text, re.IGNORECASE):
                potential_spans.append((match.start(), match.end(), kw))
                logger.debug(f"  找到[特殊]匹配: '{match.group(0)}' (规则: '{kw}') at [{match.start()}:{match.end()}]")
        except re.error as e:
            logger.error(f"  正则表达式错误 (特殊)，关键词: '{kw}', 错误: {e}")
    
    if not potential_spans:
        logger.debug("  段落中未找到任何待高亮关键词。")
        return

    # 步骤 2: 查找所有需要排除的区域
    exclude_spans = []
    for kw in exclude_keywords:
        try:
            pattern = create_pattern_for_keyword(kw, 'exclude')
            for match in re.finditer(pattern, paragraph_text, re.IGNORECASE):
                exclude_spans.append((match.start(), match.end(), kw))
                logger.debug(f"  发现[排除]区域: '{match.group(0)}' at [{match.start()}:{match.end()}]")
        except re.error as e:
            logger.error(f"  正则表达式错误 (排除)，关键词: '{kw}', 错误: {e}")

    # 步骤 3: 应用排除规则
    final_spans = []
    for p_start, p_end, p_kw in potential_spans:
        is_excluded = False
        for e_start, e_end, e_kw in exclude_spans:
            # 如果高亮区域完全被排除区域包含，则排除
            if p_start >= e_start and p_end <= e_end:
                is_excluded = True
                logger.debug(f"  排除: '{paragraph_text[p_start:p_end]}' at [{p_start}:{p_end}] 因为它被 '{e_kw}' at [{e_start}:{e_end}] 包含。")
                break
        if not is_excluded:
            final_spans.append((p_start, p_end))

    if not final_spans:
        logger.debug("  所有待高亮项均被排除。")
        return

    # 步骤 4: 合并重叠或相邻的高亮区域
    if not final_spans: return
    final_spans.sort(key=lambda x: x[0])
    
    merged_spans = []
    current_start, current_end = final_spans[0]
    for next_start, next_end in final_spans[1:]:
        if next_start < current_end: # 使用 < 而不是 <= 防止仅相邻的区域合并
            current_end = max(current_end, next_end)
        else:
            merged_spans.append((current_start, current_end))
            current_start, current_end = next_start, next_end
    merged_spans.append((current_start, current_end))

    logger.info(f"  段落处理完成，最终确定 {len(merged_spans)} 个高亮区域。")
    logger.debug(f"  最终高亮区域: {merged_spans}")

    # 步骤 5: 重建段落 (此部分逻辑无需修改，它依赖于正确的spans)
    original_runs = list(paragraph.runs)
    paragraph.clear()

    last_pos = 0 
    
    # 辅助函数，用于从原始runs中提取文本并添加到新段落
    def add_text_from_runs(start_char, end_char, is_highlight):
        text_len_to_add = end_char - start_char
        if text_len_to_add <= 0: return

        # 定位起始 run
        char_cursor = 0
        run_index = 0
        
        # 找到包含 start_char 的 run
        for i, run in enumerate(original_runs):
            run_len = len(run.text)
            if char_cursor + run_len > start_char:
                run_index = i
                break
            char_cursor += run_len
        
        # 从定位到的 run 开始添加文本
        offset_in_run = start_char - char_cursor

        while text_len_to_add > 0 and run_index < len(original_runs):
            run = original_runs[run_index]
            text_from_run = run.text[offset_in_run:]
            
            part_to_add = text_from_run
            if len(part_to_add) > text_len_to_add:
                part_to_add = part_to_add[:text_len_to_add]
            
            if part_to_add:
                new_run = paragraph.add_run(part_to_add)
                copy_run_format(run, new_run)
                if is_highlight:
                    new_run.font.highlight_color = highlight_color
            
            text_len_to_add -= len(part_to_add)
            
            # 移动到下一个 run
            run_index += 1
            offset_in_run = 0 # 新的run从头开始
    
    for h_start, h_end in merged_spans:
        # 添加非高亮部分
        add_text_from_runs(last_pos, h_start, is_highlight=False)
        # 添加高亮部分
        add_text_from_runs(h_start, h_end, is_highlight=True)
        last_pos = h_end

    # 添加最后剩余的非高亮部分
    add_text_from_runs(last_pos, len(paragraph_text), is_highlight=False)


# --- 主处理函数 ---
def highlight_keywords_in_doc(doc_path, highlight_keywords, exact_keywords, exclude_keywords, output_path, highlight_color):
    """在 Word 文档（包括正文和表格）中高亮关键词"""
    logger.info(f"\n开始处理 Word 文档: {doc_path}")
    try:
        document = docx.Document(doc_path)
        logger.info("Word 文档加载成功。")
    except Exception as e:
        logger.error(f"错误: 无法加载 Word 文档 {doc_path}: {e}")
        return

    # --- 处理主文档段落 ---
    logger.info("--- 正在处理主文档段落 ---")
    total_paras = len(document.paragraphs)
    for i, p in enumerate(document.paragraphs):
        logger.info(f"处理主段落 {i + 1}/{total_paras}")
        process_and_highlight_paragraph(p, highlight_keywords, exact_keywords, exclude_keywords, highlight_color)
    logger.info("--- 主文档段落处理完成 ---")

    # --- 处理表格中的段落 ---
    logger.info("--- 正在处理表格内容 ---")
    if document.tables:
        logger.info(f"文档中找到 {len(document.tables)} 个表格。")
        for i, table in enumerate(document.tables):
            logger.info(f"  处理表格 {i + 1}/{len(document.tables)}")
            for row in table.rows:
                for cell in row.cells:
                    for p_in_cell in cell.paragraphs:
                        process_and_highlight_paragraph(p_in_cell, highlight_keywords, exact_keywords, exclude_keywords, highlight_color)
    else:
        logger.info("  文档中未发现表格。")
    logger.info("--- 表格内容处理完成 ---")

    # --- 保存文档 ---
    try:
        logger.info(f"准备保存高亮后的文档到: {output_path}")
        document.save(output_path)
        logger.info(f"文档保存成功！输出文件位于: {os.path.abspath(output_path)}")
    except PermissionError:
        logger.error(f"错误: 保存文档失败 {output_path}。文件可能已被其他程序打开或无写入权限。")
    except Exception as e:
        logger.error(f"错误: 保存文档时发生未知错误 {output_path}: {e}", exc_info=True)


# --- 主程序入口 ---
if __name__ == "__main__":
    logger.info("--- 开始执行关键词高亮脚本 ---")
    
    # 自动初始化默认文件夹和文件
    ensure_default_keywords_files()

    if not os.path.exists(WORD_DOCUMENT_PATH):
        logger.error(f"错误: 目标Word文档未找到: {WORD_DOCUMENT_PATH}")
    else:
        highlight_keywords = load_keywords(HIGHLIGHT_KEYWORDS_PATH)
        exact_keywords = load_keywords(EXACT_KEYWORDS_PATH)
        exclude_keywords = load_keywords(EXCLUDE_KEYWORDS_PATH)

        if not highlight_keywords and not exact_keywords:
            logger.warning("警告: '高亮关键词' 和 '特殊关键词' 列表都为空，脚本将不会高亮任何内容。")
        else:
            highlight_keywords_in_doc(
                WORD_DOCUMENT_PATH,
                highlight_keywords,
                exact_keywords,
                exclude_keywords,
                OUTPUT_DOCUMENT_PATH,
                HIGHLIGHT_COLOR
            )

    logger.info("--- 脚本执行结束 ---")